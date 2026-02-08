"""
DOCX Content Extractor
======================

Extracts all content from a DOCX file in proper order, including tables.

This script solves the common problem where `doc.paragraphs` only returns 
paragraph elements and completely skips tables. By iterating through 
`doc.element.body`, we capture everything in the correct sequence.

Usage:
------
    from docx_extractor import extract_docx_content
    
    content = extract_docx_content('questionnaire.docx')
    
    for item in content:
        if item['type'] == 'paragraph':
            print(f"Paragraph: {item['text']}")
        elif item['type'] == 'table':
            print(f"Table with {len(item['data'])} rows")

Returns:
--------
List of dictionaries with structure:
    {
        'type': 'paragraph' | 'table',
        'index': int,           # Position in document
        'text': str,            # For paragraphs
        'data': list[list],     # For tables (2D array of cells)
    }
"""

from docx import Document
from typing import List, Dict, Any


def extract_docx_content(filepath: str) -> List[Dict[str, Any]]:
    """
    Extract all content from a DOCX file in proper order.
    
    Args:
        filepath: Path to the .docx file
        
    Returns:
        List of content items (paragraphs and tables) in document order
    """
    doc = Document(filepath)
    content = []
    
    # Iterate through actual document structure
    for i, element in enumerate(doc.element.body):
        
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:  # Only include non-empty paragraphs
                        content.append({
                            'type': 'paragraph',
                            'index': i,
                            'text': text
                        })
                    break
        
        # Handle tables
        elif element.tag.endswith('tbl'):
            # Find the matching table object
            for table in doc.tables:
                if table._element == element:
                    table_data = []
                    
                    # Extract all cells from all rows
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        table_data.append(cells)
                    
                    content.append({
                        'type': 'table',
                        'index': i,
                        'data': table_data
                    })
                    break
    
    return content


def validate_extraction(doc: Document, content: List[Dict[str, Any]]) -> Dict[str, int]:
    """
    Validate that extraction captured all content.
    
    Args:
        doc: Document object
        content: Extracted content list
        
    Returns:
        Dictionary with counts for validation
    """
    paragraph_count = len([c for c in content if c['type'] == 'paragraph'])
    table_count = len([c for c in content if c['type'] == 'table'])
    total_elements = len(doc.element.body)
    
    return {
        'paragraphs': paragraph_count,
        'tables': table_count,
        'total_extracted': len(content),
        'total_elements': total_elements
    }


def print_content_summary(content: List[Dict[str, Any]]) -> None:
    """
    Print a summary of extracted content.
    
    Args:
        content: Extracted content list
    """
    print(f"\n{'='*60}")
    print("DOCX Content Extraction Summary")
    print(f"{'='*60}\n")
    
    paragraph_count = 0
    table_count = 0
    
    for item in content:
        if item['type'] == 'paragraph':
            paragraph_count += 1
            text_preview = item['text'][:60] + '...' if len(item['text']) > 60 else item['text']
            print(f"[P{paragraph_count:02d}] {text_preview}")
            
        elif item['type'] == 'table':
            table_count += 1
            rows = len(item['data'])
            cols = len(item['data'][0]) if item['data'] else 0
            print(f"[T{table_count:02d}] Table: {rows} rows × {cols} columns")
    
    print(f"\n{'='*60}")
    print(f"Total: {paragraph_count} paragraphs, {table_count} tables")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <path_to_docx_file>")
        sys.exit(1)
    
    filepath = sys.argv[1]
    
    try:
        # Extract content
        content = extract_docx_content(filepath)
        
        # Print summary
        print_content_summary(content)
        
        # Validate extraction
        doc = Document(filepath)
        validation = validate_extraction(doc, content)
        
        print("Validation:")
        print(f"  - Paragraphs extracted: {validation['paragraphs']}")
        print(f"  - Tables extracted: {validation['tables']}")
        print(f"  - Total items: {validation['total_extracted']}")
        print(f"  - Total document elements: {validation['total_elements']}")
        
        if validation['total_extracted'] < validation['total_elements']:
            print("\n⚠️  Warning: Some elements may not have been extracted!")
        else:
            print("\n✓ All content extracted successfully!")
            
    except Exception as e:
        print(f"\n❌ Error: {e}")
        sys.exit(1)
